"""Document generation module for the Demo Brief Generator.

Creates professional Word documents (.docx) with structured prospect
and company research data.
"""

import os
import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import config


def _set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="0066CC"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t>{_escape_xml(text)}</w:t>'
        f'</w:r>'
        f'</w:hyperlink>'
    )
    paragraph._element.append(hyperlink)
    return hyperlink


def _escape_xml(text):
    """Escape special XML characters."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _set_font(run, name=None, size=None, bold=None, italic=None, color=None):
    """Configure font properties on a run."""
    if name:
        run.font.name = name
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_bullet(doc, text, level=0, bold=False, italic=False, color=None, font_size=None):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.style = doc.styles["List Bullet 2"] if level == 1 else doc.styles["List Bullet 3"]
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))

    run = p.add_run(text)
    _set_font(
        run,
        name=config.FONT_NAME,
        size=font_size or config.FONT_SIZE_BODY,
        bold=bold,
        italic=italic,
        color=color,
    )
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    return p


def _add_bullet_with_link(doc, text_before, link_text, url, text_after="", level=0):
    """Add a bullet point with an embedded hyperlink."""
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        try:
            p.style = doc.styles["List Bullet 2"]
        except KeyError:
            pass
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))

    if text_before:
        run = p.add_run(text_before)
        _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_BODY)

    _add_hyperlink(p, link_text, url)

    if text_after:
        run = p.add_run(text_after)
        _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_BODY)

    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    return p


def _add_section_heading(doc, text, level=1):
    """Add a section heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: config.FONT_SIZE_HEADING, 2: config.FONT_SIZE_SUBHEADING}
    _set_font(
        run,
        name=config.FONT_NAME,
        size=sizes.get(level, config.FONT_SIZE_BODY),
        bold=True,
        color=config.COLOR_PRIMARY,
    )
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_sub_header(doc, text):
    """Add a sub-header within a section."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(
        run,
        name=config.FONT_NAME,
        size=config.FONT_SIZE_SUBHEADING,
        bold=True,
        color=config.COLOR_PRIMARY,
    )
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_note(doc, text, color=None):
    """Add a note paragraph (italic, colored)."""
    p = doc.add_paragraph()
    run = p.add_run(f"Note: {text}")
    _set_font(
        run,
        name=config.FONT_NAME,
        size=config.FONT_SIZE_BODY,
        italic=True,
        color=color or config.COLOR_GRAY,
    )
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.27)
    return p


def _load_logo(logo_path):
    """Load a logo image, converting SVG to PNG if necessary. Returns (stream, width_inches)."""
    if not logo_path or not os.path.exists(logo_path):
        return None, None

    _, ext = os.path.splitext(logo_path.lower())
    if ext == ".svg":
        try:
            import cairosvg
            png_data = cairosvg.svg2png(url=logo_path, output_width=200)
            return io.BytesIO(png_data), Inches(1.5)
        except Exception as e:
            print(f"  [!] Failed to convert SVG logo: {e}")
            return None, None
    elif ext in (".png", ".jpg", ".jpeg"):
        return open(logo_path, "rb"), Inches(1.5)
    else:
        print(f"  [!] Unsupported logo format: {ext}")
        return None, None


def _not_found(value):
    """Check if a value is 'Not found' or empty."""
    if isinstance(value, str):
        return value in ("Not found", "", "N/A")
    if isinstance(value, list):
        return len(value) == 0
    return value is None


def generate_docx(prospect_data, company_data, sdr_name, logo_path=None, output_path=None):
    """Generate a professional Word document from research data.

    Args:
        prospect_data: Dict with prospect research results
        company_data: Dict with company research results
        sdr_name: Name of the SDR
        logo_path: Optional path to company logo file
        output_path: Optional output file path

    Returns:
        Path to the generated .docx file
    """
    prospect_name = prospect_data["name"]
    company_name = company_data["name"]

    if output_path is None:
        safe_prospect = prospect_name.replace(" ", "_")
        safe_company = company_name.replace(" ", "_")
        output_path = os.path.join(
            config.OUTPUT_DIR,
            f"{safe_prospect}_-_{safe_company}_-_Demo_Brief.docx",
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # Configure default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = config.FONT_NAME
    font.size = Pt(config.FONT_SIZE_BODY)

    # Configure page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ==========================================
    # HEADER (Table Layout)
    # ==========================================
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(4.5)
    header_table.columns[1].width = Inches(2)

    # Left cell: Title info
    left_cell = header_table.cell(0, 0)
    left_cell.vertical_alignment = 1  # CENTER

    # "Demo Brief" title
    p = left_cell.paragraphs[0]
    run = p.add_run("Demo Brief")
    _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_TITLE, bold=True, color=config.COLOR_PRIMARY)
    p.paragraph_format.space_after = Pt(2)

    # Prospect name
    p = left_cell.add_paragraph()
    run = p.add_run(prospect_name)
    _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_HEADING, bold=True, color=config.COLOR_PRIMARY)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    # SDR name
    p = left_cell.add_paragraph()
    run = p.add_run(f"SDR: {sdr_name}")
    _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_BODY, italic=True, color=config.COLOR_GRAY)
    p.paragraph_format.space_before = Pt(0)

    # Right cell: Logo
    right_cell = header_table.cell(0, 1)
    right_cell.vertical_alignment = 1  # CENTER

    logo_stream, logo_width = _load_logo(logo_path)
    if logo_stream:
        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_stream, width=logo_width)
        if hasattr(logo_stream, 'close'):
            logo_stream.close()

        # Tagline below logo
        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{company_name}")
        _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_SMALL, italic=True, color=config.COLOR_GRAY)

    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0"/>'
                '<w:left w:val="none" w:sz="0" w:space="0"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                '<w:right w:val="none" w:sz="0" w:space="0"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    # Divider line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # Add bottom border to paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ==========================================
    # INFORMATION TABLE
    # ==========================================
    info_fields = [
        ("Company", company_name, company_data.get("website")),
        ("Name", prospect_name, prospect_data.get("linkedin_url")),
        ("Prospect Location", prospect_data.get("location", "Not found"), None),
        ("Role", _format_role(prospect_data), None),
        ("Company Size", company_data.get("size", "Not found"), None),
        ("Industry", company_data.get("industry", "Not found"), None),
        ("ATS", company_data.get("ats", "Not found"), None),
        ("Current Open Remote Jobs", company_data.get("open_remote_jobs", "Not found"), None),
        ("Headquarters", company_data.get("headquarters", "Not found"), None),
    ]

    info_table = doc.add_table(rows=len(info_fields), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.columns[0].width = Inches(2.2)
    info_table.columns[1].width = Inches(4.3)

    for i, (label, value, link) in enumerate(info_fields):
        # Label cell
        label_cell = info_table.cell(i, 0)
        _set_cell_shading(label_cell, "FFF4E6")
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_BODY, bold=True, color=config.COLOR_PRIMARY)

        # Value cell
        value_cell = info_table.cell(i, 1)
        p = value_cell.paragraphs[0]

        is_missing = _not_found(value)
        if link and not _not_found(link) and link != "Not found":
            _add_hyperlink(p, str(value), link)
        else:
            run = p.add_run(str(value))
            color = config.COLOR_RED if is_missing else None
            _set_font(run, name=config.FONT_NAME, size=config.FONT_SIZE_BODY, color=color)

    # Style table borders
    _style_table_borders(info_table)

    doc.add_paragraph()  # spacer

    # ==========================================
    # SECTION 1: About [Prospect Name]
    # ==========================================
    _add_section_heading(doc, f"\U0001f3af About {prospect_name}")

    # Role
    role_text = prospect_data.get("title", "Not found")
    if _not_found(role_text):
        _add_bullet(doc, "Role: Not found", color=config.COLOR_RED, italic=True)
    else:
        _add_bullet(doc, f"Role: {role_text}")

    # Team
    team = prospect_data.get("team", "Not found")
    if _not_found(team):
        _add_bullet(doc, "Team: Not found", color=config.COLOR_RED, italic=True)
    else:
        _add_bullet(doc, f"Team: {team}")

    # Certifications
    certs = prospect_data.get("certifications", [])
    if certs:
        _add_bullet(doc, f"Certifications: {', '.join(certs)}")
    else:
        _add_bullet(doc, "Certifications: Not found", color=config.COLOR_RED, italic=True)

    # Work History
    _add_sub_header(doc, "Work History:")
    work_history = prospect_data.get("work_history", [])
    if work_history:
        for entry in work_history[:8]:
            _add_bullet(doc, entry, level=0)
    else:
        _add_note(doc, "Work history could not be determined from public sources.", config.COLOR_RED)

    # Notable Achievements
    _add_sub_header(doc, "Notable Achievements (Identity Security & Deepfakes):")
    achievements = prospect_data.get("achievements", [])
    if achievements:
        for ach in achievements[:5]:
            _add_bullet(doc, ach, level=0)
    else:
        _add_note(
            doc,
            "No specific achievements in identity security or deepfakes found in public sources.",
            config.COLOR_RED,
        )

    # Published Content
    _add_sub_header(doc, "Published Content & Thought Leadership:")
    published = prospect_data.get("published_content", [])
    if published:
        for content in published[:8]:
            date_str = f" ({content['date']})" if content.get("date") else ""
            type_str = "[Talk] " if content.get("type") == "talk" else ""
            if content.get("url"):
                _add_bullet_with_link(
                    doc,
                    type_str,
                    content["title"],
                    content["url"],
                    text_after=date_str,
                )
            else:
                _add_bullet(doc, f"{type_str}{content['title']}{date_str}")
    else:
        _add_note(doc, "No published content found in public sources.", config.COLOR_GRAY)

    # ==========================================
    # SECTION 2: Company Overview
    # ==========================================
    _add_section_heading(doc, f"\U0001f9e9 Company Overview")

    # Company basics
    founded = company_data.get("founded", "Not found")
    ticker = company_data.get("ticker", "Not found")
    hq = company_data.get("headquarters", "Not found")

    basics_parts = []
    basics_parts.append(f"Founded: {founded}" if not _not_found(founded) else "Founded: Not found")
    if not _not_found(ticker):
        basics_parts.append(f"Public: {ticker}")
    basics_parts.append(f"HQ: {hq}" if not _not_found(hq) else "HQ: Not found")
    _add_bullet(doc, " | ".join(basics_parts))

    # Product
    product = company_data.get("product_description", "Not found")
    if _not_found(product):
        _add_bullet(doc, "Product: Not found", color=config.COLOR_RED, italic=True)
    else:
        _add_bullet(doc, f"Product: {product}")

    # Customers
    customers = company_data.get("customers", "Not found")
    if _not_found(customers):
        _add_bullet(doc, "Customers: Not found", color=config.COLOR_RED, italic=True)
    else:
        _add_bullet(doc, f"Customers: {customers}")

    # Culture
    culture = company_data.get("culture", "Not found")
    if _not_found(culture):
        _add_bullet(doc, "Culture: Not found", color=config.COLOR_RED, italic=True)
    else:
        _add_bullet(doc, f"Culture: {culture}")

    # --- Hiring Growth & Recruitment Infrastructure ---
    _add_sub_header(doc, "Hiring Growth & Recruitment Infrastructure:")

    emp = company_data.get("employee_count", "Not found")
    growth = company_data.get("growth", "Not found")
    emp_text = f"Employees: {emp}"
    if not _not_found(growth):
        emp_text += f" | Growth: {growth}"
    _add_bullet(doc, emp_text, color=config.COLOR_RED if _not_found(emp) else None)

    hiring = company_data.get("hiring_activity", "Not found")
    _add_bullet(
        doc,
        f"Hiring Activity: {hiring}",
        color=config.COLOR_RED if _not_found(hiring) else None,
        italic=_not_found(hiring),
    )

    team_struct = company_data.get("team_structure", "Not found")
    _add_bullet(
        doc,
        f"Team Structure: {team_struct}",
        color=config.COLOR_RED if _not_found(team_struct) else None,
        italic=_not_found(team_struct),
    )

    ats = company_data.get("ats", "Not found")
    _add_bullet(
        doc,
        f"ATS: {ats}",
        color=config.COLOR_RED if _not_found(ats) else None,
        italic=_not_found(ats),
    )

    open_roles = company_data.get("open_remote_jobs", "Not found")
    _add_bullet(
        doc,
        f"Open Roles: {open_roles} positions" if not _not_found(open_roles) else "Open Roles: Not found",
        color=config.COLOR_RED if _not_found(open_roles) else None,
        italic=_not_found(open_roles),
    )

    # --- Identity & Access Management Tools ---
    _add_sub_header(doc, "Identity & Access Management Tools:")
    id_tools = company_data.get("identity_tools", [])
    if id_tools:
        for tool in id_tools:
            _add_bullet(doc, f"{tool['name']} - {tool['description']}")
    else:
        _add_note(doc, "No specific identity tools identified from public sources.", config.COLOR_RED)

    compliance = company_data.get("compliance", "Not found")
    _add_bullet(
        doc,
        f"Compliance: {compliance}",
        color=config.COLOR_RED if _not_found(compliance) else None,
        italic=_not_found(compliance),
    )

    # --- Recent Security Incidents ---
    _add_sub_header(doc, "Recent Security Incidents:")
    incidents = company_data.get("security_incidents", [])
    if incidents:
        for inc in incidents[:5]:
            _add_bullet(doc, f"{inc['date']}: {inc['title']}", bold=True)
            # Sub-bullet with details
            details = inc.get("details", "")
            if details:
                _add_bullet(doc, details, level=1)
    else:
        _add_bullet(doc, "No recent security incidents found in public sources.", italic=True)

    # Hiring security notes
    hiring_notes = company_data.get("hiring_security_notes", [])
    if hiring_notes:
        _add_bullet(doc, "Highlights:", bold=True)
        for note in hiring_notes[:3]:
            _add_bullet(doc, note, level=1)
    else:
        _add_bullet(
            doc,
            "No hiring/recruitment-related security incidents found.",
            italic=True,
            color=config.COLOR_GRAY,
        )

    # Save document
    doc.save(output_path)
    return output_path


def _format_role(prospect_data):
    """Format role string with tenure info."""
    title = prospect_data.get("title", "Not found")
    tenure = prospect_data.get("tenure", "")
    if _not_found(title):
        return "Not found"
    if tenure and not _not_found(tenure):
        return f"{title} - {tenure}"
    return title


def _style_table_borders(table):
    """Apply light borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
